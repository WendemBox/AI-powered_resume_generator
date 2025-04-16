from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from ai_generator import generate_resume
import os
from dotenv import load_dotenv
import logging
import tempfile
from docx import Document
import markdown
import pdfkit
from datetime import datetime
from docx import Document

# Загрузка переменных окружения
load_dotenv()

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('flask_app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, supports_credentials=True)

# Конфигурация
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# Конфигурация для pdfkit
PDFKIT_CONFIG = {
    'page-size': 'A4',
    'margin-top': '0.75in',
    'margin-right': '0.75in',
    'margin-bottom': '0.75in',
    'margin-left': '0.75in',
    'encoding': 'UTF-8',
    'quiet': ''
}

@app.route('/api/generate-resume', methods=['POST'])
def generate_resume_endpoint():
    """Генерация текста резюме с помощью Yandex GPT"""
    try:
        logger.info("Received request to /api/generate-resume")
        
        if not request.is_json:
            logger.warning("Request does not contain JSON data")
            return jsonify({
                'success': False,
                'message': 'Ожидается JSON в теле запроса'
            }), 400
            
        data = request.json
        logger.info(f"Received data for resume generation")
        
        user_data = data.get('user_data', {})
        resume_type = data.get('resume_type', 'standard')
        
        # Валидация данных
        if not user_data.get('name'):
            return jsonify({
                'success': False,
                'message': 'Поле "ФИО" обязательно для заполнения'
            }), 400

        # Генерация резюме
        logger.info(f"Generating {resume_type} resume")
        resume_content = generate_resume(user_data, resume_type)
        
        logger.info("Resume generated successfully")
        return jsonify({
            'success': True,
            'resume': resume_content,
            'message': 'Резюме успешно сгенерировано'
        })
    except Exception as e:
        logger.error(f"Error generating resume: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Ошибка при генерации резюме: {str(e)}'
        }), 500

@app.route('/api/download-resume', methods=['POST'])
def download_resume():
    """Генерация и скачивание резюме в выбранном формате"""
    try:
        logger.info("Received request to /api/download-resume")
        
        if not request.is_json:
            return jsonify({
                'success': False,
                'message': 'Ожидается JSON в теле запроса'
            }), 400
            
        data = request.json
        resume_content = data.get('resume_content', '')
        format_type = data.get('format', 'pdf').lower()
        file_name = data.get('file_name', f'resume_{datetime.now().strftime("%Y%m%d")}')
        
        if not resume_content:
            return jsonify({
                'success': False,
                'message': 'Отсутствует содержимое резюме'
            }), 400
            
        if format_type not in ['pdf', 'docx']:
            return jsonify({
                'success': False,
                'message': 'Неподдерживаемый формат файла'
            }), 400

        # Создаем временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{format_type}') as temp_file:
            temp_path = temp_file.name
            
            if format_type == 'pdf':
                # Конвертация Markdown в HTML
                html_content = markdown.markdown(resume_content)
                # Добавляем базовые стили
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
                        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; }}
                        h2 {{ color: #2980b9; }}
                        ul {{ margin-left: 20px; }}
                        li {{ margin-bottom: 5px; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                pdfkit.from_string(html_content, temp_path, options=PDFKIT_CONFIG)
                
            elif format_type == 'docx':
                doc = Document()
                
                # Парсим Markdown и добавляем в DOCX
                current_paragraph = None
                for line in resume_content.split('\n'):
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('- '):
                        if current_paragraph is None:
                            current_paragraph = doc.add_paragraph()
                        current_paragraph.add_run('• ' + line[2:] + '\n')
                    else:
                        if line.strip():
                            doc.add_paragraph(line)
                        current_paragraph = None
                
                doc.save(temp_path)
        
        logger.info(f"File generated successfully: {temp_path}")
        
        # Отправка файла клиенту
        return send_file(
            temp_path,
            as_attachment=True,
            download_name=f"{file_name}.{format_type}",
            mimetype=f"application/{format_type}"
        )
        
    except Exception as e:
        logger.error(f"Error generating download file: {str(e)}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Ошибка при генерации файла: {str(e)}'
        }), 500
    finally:
        # Удаление временного файла после отправки
        try:
            if 'temp_path' in locals() and os.path.exists(temp_path):
                os.unlink(temp_path)
        except Exception as e:
            logger.error(f"Error deleting temp file: {str(e)}")

@app.route('/api/health', methods=['GET'])
def health_check():
    """Проверка работоспособности API"""
    return jsonify({
        'status': 'OK',
        'message': 'Resume Generator API is running',
        'version': '1.0.0',
        'supported_formats': ['pdf', 'docx']
    })

if __name__ == '__main__':
    # Создаем папку для загрузок, если её нет
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    
    # Запускаем сервер
    app.run(debug=os.getenv('FLASK_DEBUG', 'false').lower() == 'true', 
           port=int(os.getenv('FLASK_PORT', 5000)),
           host=os.getenv('FLASK_HOST', '0.0.0.0'))